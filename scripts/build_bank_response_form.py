"""Generate sendable Cornerstone bank response form (DOCX + PDF) for 3887 Seneca / Luxe Loft 716."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

OUT_DIR = Path(r"C:\Users\17274\ME\3887-Seneca\docs\bank")
DOCX_PATH = OUT_DIR / "Cornerstone_Bank_Response_Form.docx"
PDF_PATH = OUT_DIR / "Cornerstone_Bank_Response_Form.pdf"


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_line(doc, text, size=16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x5C, 0x24, 0x30)
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        for para in table.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
    doc.add_paragraph()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("COMMERCIAL LOAN RESPONSE FORM")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x5C, 0x24, 0x30)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Luxe Loft 716 — 3887 Seneca Street\nPrepared for Cornerstone Community Federal Credit Union\nAttention: Michael Woods · July 2026")
    sr.font.size = Pt(11)

    note = doc.add_paragraph()
    nr = note.add_run(
        "This form answers the underwriting items needed to complete the financing request for 3887 Seneca. "
        "Attach with: (1) IJB Commercial Financial Statement Package, (2) Sale Contract, (3) Luxe Loft 716 Presentation."
    )
    nr.italic = True
    nr.font.size = Pt(9)

    add_heading_line(doc, "1. Borrower & Contact", 13)
    add_kv_table(doc, [
        ("Borrower / Purchaser", "IJB Enterprises Inc."),
        ("Project / Brand", "Luxe Loft 716"),
        ("Primary Contact", "Tiffany Durilla"),
        ("Email", "durillaprop@gmail.com"),
        ("Phone", "716-421-1210"),
        ("Broker", "Steiner & Blotnik"),
        ("Seller Counsel", "Bill Miller Law / Bill Miller"),
        ("Related entities (package)", "549 Indian Church LLC; Durilla Development of WNY LLC"),
    ])

    add_heading_line(doc, "2. Property", 13)
    add_kv_table(doc, [
        ("Address", "3887 Seneca Street, West Seneca, NY 14224-3489"),
        ("SBL", "146800-143-070-0006-005-000"),
        ("Lot size", "Approximately 64 × 173"),
        ("Current / intended use", "Multi-use commercial — two-family dwelling + commercial salon suites"),
        ("Prior use", "Former dental / professional office"),
        ("1st-floor SF (measured)", "TO BE PROVIDED (pending measurement)"),
    ])

    add_heading_line(doc, "3. Purchase Contract Terms", 13)
    add_kv_table(doc, [
        ("Contract date", "June 16, 2026"),
        ("Authentisign ID", "96DFCCA5-C06A-F111-8FCA-002248359474"),
        ("Seller", "Seneca Street 3887"),
        ("Purchase price", "$500,000.00"),
        ("Seller concession", "$50,000.00"),
        ("Net consideration", "$450,000.00"),
        ("Deposit", "$10,000.00 (Kreag Ferullo / Steiner & Blotnik escrow)"),
        ("Seller repair (contract)", "Exterior main electric service line repair/upgrade at Seller's sole expense (Town of West Seneca Electrical Inspector: Steve Carmina)"),
    ])

    add_heading_line(doc, "4. Business Purpose of Loan", 13)
    doc.add_paragraph(
        "IJB Enterprises Inc. is acquiring 3887 Seneca Street to operate Luxe Loft 716 — a luxury salon and wellness "
        "suite business for independent beauty and wellness professionals — and to continue leasing two residential "
        "apartments. The first floor will be configured as nine private lockable suites. Revenue comes from multiple "
        "suite leases plus residential rents, increasing income per square foot and reducing single-tenant vacancy risk."
    )

    add_heading_line(doc, "5. Stabilized Income & Debt Coverage (Luxe Loft 716 Underwriting)", 13)
    add_kv_table(doc, [
        ("9 salon suites (annual)", "$108,000"),
        ("Apartment 1 (annual)", "$13,200"),
        ("Apartment 2 (annual)", "$15,600"),
        ("Residential combined", "$2,400 / month · $28,800 / year"),
        ("Gross Potential Income (GPI)", "$136,800"),
        ("Operating expenses (implied)", "$40,880"),
        ("Net Operating Income (NOI)", "$95,920"),
        ("Debt service (underwritten)", "$31,944 / year"),
        ("DSCR", "3.00"),
    ])
    p = doc.add_paragraph()
    r = p.add_run(
        "Note: Residential underwriting baseline is $2,400/mo from the Luxe Loft 716 presentation. "
        "Lease copies will confirm final rents. Debt service will be updated when Cornerstone quotes final loan terms."
    )
    r.italic = True
    r.font.size = Pt(9)

    add_heading_line(doc, "6. Sources & Uses", 13)
    add_kv_table(doc, [
        ("USE — Net purchase", "$450,000"),
        ("USE — Closing costs", "TO BE FINALIZED"),
        ("USE — Suite conversion CapEx", "TO BE FINALIZED (contractor bids)"),
        ("USE — Reserves / working capital", "TO BE FINALIZED"),
        ("SOURCE — Cornerstone loan", "TO BE FINALIZED (loan ask after CapEx)"),
        ("SOURCE — Borrower equity (incl. $10,000 deposit)", "TO BE FINALIZED"),
    ])

    add_heading_line(doc, "7. Conversion / Build-Out Plan", 13)
    for item in [
        "Measure first-floor square footage; recover prior dental architect-stamped plans.",
        "Finalize nine-suite layout with lobby/common area; prefer demountable partitions (STC ≥ 40).",
        "Provide per-suite power, Wi-Fi, and wet/vent stubs as required for beauty/wellness uses.",
        "Market to independent operators; lease-up to stabilize suite income at underwritten levels.",
        "Confirm Seller completes exterior electric service repair/inspection before or at closing.",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading_line(doc, "8. Credit Strength Summary", 13)
    for item in [
        "Diversified rent roll: nine suites + two apartments vs. a single commercial tenant.",
        "Existing professional office infrastructure (plumbing, HVAC, parking, ADA) lowers conversion risk.",
        "Prior architect-designed dental conversion cited in Luxe Loft 716 presentation.",
        "Underwritten DSCR of 3.00 on presentation assumptions.",
        "Related operating history provided via 549 Indian Church LLC in the financial package.",
        "Buffalo salon-suite operating experience with independent operators / micro-suites.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_line(doc, "9. Loan Request (complete before send)", 13)
    add_kv_table(doc, [
        ("Requested loan amount", "$________________"),
        ("Requested term / amortization", "________________"),
        ("Requested rate structure", "________________"),
        ("Requested LTV", "________________ %"),
        ("Guarantor(s)", "________________"),
    ])

    add_heading_line(doc, "10. Attachments Checklist", 13)
    for item in [
        "IJB Enterprises Commercial Financial Statement Package (July 2026)",
        "Sale Contract — 3887 Seneca (signed / Authentisign)",
        "Luxe Loft 716 Presentation (prepared for Michael Woods)",
        "Personal Financial Statement / tax returns (guarantor) — attach when ready",
        "Entity formation documents (IJB Enterprises Inc.) — attach when ready",
        "Residential leases / rent roll — attach when ready",
        "CapEx bids / contractor estimates — attach when ready",
        "Measured floor plan / dental architect plans — attach when ready",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_line(doc, "11. Certification", 13)
    doc.add_paragraph(
        "The undersigned certifies that the information in this response form is true and correct to the best of "
        "their knowledge as of the date signed, and authorizes Cornerstone Community Federal Credit Union to verify "
        "the information provided."
    )
    doc.add_paragraph()
    add_kv_table(doc, [
        ("Authorized signer (print)", "Tiffany Durilla"),
        ("Title / entity", "IJB Enterprises Inc."),
        ("Signature", "________________________________"),
        ("Date", "________________"),
    ])

    footer = doc.add_paragraph()
    fr = footer.add_run(
        "\nInternal file location: C:\\Users\\17274\\ME\\3887-Seneca\\docs\\bank\\Cornerstone_Bank_Response_Form.docx"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(DOCX_PATH)
    print("Wrote", DOCX_PATH)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleRose", parent=styles["Title"], fontSize=16, textColor=colors.HexColor("#5c2430"), spaceAfter=6))
    styles.add(ParagraphStyle(name="HRose", parent=styles["Heading2"], fontSize=11, textColor=colors.HexColor("#5c2430"), spaceBefore=10, spaceAfter=4))
    styles.add(ParagraphStyle(name="BodySmall", parent=styles["Normal"], fontSize=9, leading=12))
    styles.add(ParagraphStyle(name="ItalicNote", parent=styles["Normal"], fontSize=8, leading=10, textColor=colors.HexColor("#555555")))

    story = []
    story.append(Paragraph("COMMERCIAL LOAN RESPONSE FORM", styles["TitleRose"]))
    story.append(Paragraph(
        "Luxe Loft 716 — 3887 Seneca Street<br/>Prepared for Cornerstone Community Federal Credit Union<br/>Attention: Michael Woods · July 2026",
        styles["BodySmall"],
    ))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#8b3a4a")))
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        "<i>Send this form with: (1) IJB Commercial Financial Statement Package, (2) Sale Contract, (3) Luxe Loft 716 Presentation.</i>",
        styles["ItalicNote"],
    ))

    def section(title, rows):
        story.append(Paragraph(title, styles["HRose"]))
        data = [[Paragraph(f"<b>{k}</b>", styles["BodySmall"]), Paragraph(str(v), styles["BodySmall"])] for k, v in rows]
        t = Table(data, colWidths=[2.3 * inch, 4.7 * inch])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f7f3ee")),
        ]))
        story.append(t)

    section("1. Borrower & Contact", [
        ("Borrower", "IJB Enterprises Inc."),
        ("Brand", "Luxe Loft 716"),
        ("Contact", "Tiffany Durilla · durillaprop@gmail.com · 716-421-1210"),
        ("Broker", "Steiner & Blotnik"),
        ("Seller counsel", "Bill Miller Law"),
    ])
    section("2. Property", [
        ("Address", "3887 Seneca Street, West Seneca, NY 14224-3489"),
        ("SBL", "146800-143-070-0006-005-000"),
        ("Lot", "~64 × 173"),
        ("Use", "Two-family + commercial salon suites (Luxe Loft 716)"),
        ("1st-floor SF", "TO BE PROVIDED"),
    ])
    section("3. Purchase Terms", [
        ("Contract date", "June 16, 2026"),
        ("Purchase price", "$500,000"),
        ("Seller concession", "$50,000"),
        ("Net consideration", "$450,000"),
        ("Deposit", "$10,000 (Kreag Ferullo / Steiner & Blotnik)"),
    ])
    story.append(Paragraph("4. Business Purpose", styles["HRose"]))
    story.append(Paragraph(
        "Acquire 3887 Seneca to operate Luxe Loft 716 (nine private salon/wellness suites for independent professionals) "
        "plus two residential apartments. Multi-tenant suite rents + residential income; reduced single-tenant vacancy risk.",
        styles["BodySmall"],
    ))
    section("5. Stabilized Underwriting (Luxe Deck)", [
        ("9 suites", "$108,000 / yr"),
        ("Residential", "$2,400 / mo · $28,800 / yr"),
        ("GPI", "$136,800"),
        ("NOI", "$95,920"),
        ("Debt service / DSCR", "$31,944 / yr · 3.00"),
    ])
    section("6. Sources & Uses (open items)", [
        ("Net purchase (use)", "$450,000"),
        ("Closing / CapEx / reserves (use)", "TO BE FINALIZED"),
        ("Cornerstone loan (source)", "TO BE FINALIZED"),
        ("Equity incl. deposit (source)", "TO BE FINALIZED"),
    ])
    story.append(Paragraph("7. Certification", styles["HRose"]))
    story.append(Paragraph(
        "Information is true and correct to the best of the undersigned’s knowledge. Cornerstone may verify.",
        styles["BodySmall"],
    ))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Authorized signer: Tiffany Durilla · IJB Enterprises Inc.", styles["BodySmall"]))
    story.append(Paragraph("Signature: _______________________________ &nbsp;&nbsp;&nbsp; Date: ______________", styles["BodySmall"]))

    SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    ).build(story)
    print("Wrote", PDF_PATH)


if __name__ == "__main__":
    build_docx()
    build_pdf()
