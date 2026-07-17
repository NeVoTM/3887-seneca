"""Build numbered Cornerstone document-request fill form + sample templates."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak, KeepTogether
)

ROOT = Path(r"C:\Users\17274\ME\3887-Seneca")
BANK = ROOT / "docs" / "bank"
TEMPL = BANK / "sample_templates"
OUT_DOCX = BANK / "Cornerstone_Document_Request_Fill_Form.docx"
OUT_PDF = BANK / "Cornerstone_Document_Request_Fill_Form.pdf"

# Numbered request items inferred from IJB package TOC + "See ..." refs + standard commercial acquisition underwriting
ITEMS = [
    {
        "num": 1,
        "title": "IJB Enterprises Inc. — Startup Statement of Income",
        "source": "Package TOC item 1",
        "status": "DRAFTED in package ($0 revenue / $0 NI)",
        "fields": [
            ("Period covered", "As of July 2026 (pre-operations)"),
            ("Revenue", "$0.00"),
            ("Organizational expenses", "$0.00"),
            ("Net income", "$0.00"),
            ("Notes", "Entity formed for acquisition of 3887 Seneca; no ops yet"),
        ],
        "sample": "01_IJB_Startup_Statement_of_Income.md",
    },
    {
        "num": 2,
        "title": "IJB Enterprises Inc. — Startup Balance Sheet",
        "source": "Package TOC item 2",
        "status": "INCOMPLETE — fill Cash / Equity before send",
        "fields": [
            ("Cash on hand / to fund at close", "$______________"),
            ("Organizational costs", "$______________"),
            ("Total assets", "$______________"),
            ("Liabilities", "$0.00 (or list)"),
            ("Owner's equity", "$______________"),
            ("Equity source (guarantor cash / gift / other)", "______________"),
        ],
        "sample": "02_IJB_Startup_Balance_Sheet.md",
    },
    {
        "num": 3,
        "title": "549 Indian Church LLC — Statement of Income",
        "source": "Package TOC item 3",
        "status": "DRAFTED (Jan 1 – Apr 30, 2026)",
        "fields": [
            ("Period", "January 1, 2026 – April 30, 2026"),
            ("2829 Niagara Street rent", "$2,333.98"),
            ("549 Unit 1 rent", "$4,200.00"),
            ("549 Unit 2 rent", "$2,800.00"),
            ("549 Unit 3 rent", "$4,400.00"),
            ("Total revenue", "$13,733.98"),
            ("Real estate taxes", "$6,200.00"),
            ("Property insurance", "$872.00"),
            ("NOI", "$6,661.98"),
            ("Attach full-year P&L if requested", "Yes / No — path: ______________"),
        ],
        "sample": "03_549_Indian_Church_Statement_of_Income.md",
    },
    {
        "num": 4,
        "title": "549 Indian Church LLC — Balance Sheet",
        "source": "Package TOC item 4 (points to supporting docs)",
        "status": "PLACEHOLDER — attach bank stmts / PFS / mortgages",
        "fields": [
            ("Cash (attach bank statements)", "See Item 7 — path: ______________"),
            ("Investment real estate value", "See Item 8 PFS — $______________"),
            ("Mortgage balances (attach statements)", "See Item 9 — total $______________"),
            ("Other liabilities", "$______________"),
            ("Owner's equity (residual)", "$______________"),
        ],
        "sample": "04_549_Indian_Church_Balance_Sheet.md",
    },
    {
        "num": 5,
        "title": "Durilla Development of WNY LLC — 561 Main Street Pro Forma SOI",
        "source": "Package TOC item 5",
        "status": "DRAFTED in package",
        "fields": [
            ("Projected gross income", "$171,196"),
            ("Projected operating expenses", "$71,359"),
            ("Projected NOI", "$99,837"),
            ("Est. value @ 8% cap", "$1,247,963"),
            ("Est. value @ 7% cap", "$1,426,244"),
            ("Supporting comps / rent roll attached?", "Yes / No — path: ______________"),
        ],
        "sample": "05_Durilla_561_Main_Pro_Forma.md",
    },
    {
        "num": 6,
        "title": "Supporting Documents — Cover Index",
        "source": "Package TOC item 6 (list was empty — this form completes it)",
        "status": "USE THIS FORM as the supporting-docs checklist",
        "fields": [
            ("Index prepared by", "______________"),
            ("Date submitted to Cornerstone", "______________"),
            ("Submitted to", "Michael Woods"),
        ],
        "sample": "06_Supporting_Documents_Cover_Index.md",
    },
    {
        "num": 7,
        "title": "Bank Statements (business + personal)",
        "source": "Referenced on 549 BS: “See Bank Statements”",
        "status": "NEEDED",
        "fields": [
            ("Business accounts (entity / months)", "______________ · last ___ months"),
            ("Personal accounts (guarantor / months)", "______________ · last ___ months"),
            ("File path(s)", "______________"),
            ("Avg collected balance (optional)", "$______________"),
        ],
        "sample": "07_Bank_Statements_Cover_Sheet.md",
    },
    {
        "num": 8,
        "title": "Personal Financial Statement (Guarantor)",
        "source": "Referenced on 549 BS: “See Personal Financial Statement”",
        "status": "NEEDED",
        "fields": [
            ("Guarantor legal name", "______________"),
            ("As-of date", "______________"),
            ("Total assets", "$______________"),
            ("Total liabilities", "$______________"),
            ("Net worth", "$______________"),
            ("Annual income (salary / other)", "$______________"),
            ("File path", "______________"),
        ],
        "sample": "08_Personal_Financial_Statement_Template.md",
    },
    {
        "num": 9,
        "title": "Mortgage Statements (related properties)",
        "source": "Referenced on 549 BS: “Per Mortgage Statements”",
        "status": "NEEDED",
        "fields": [
            ("Property 1 address / lender / balance", "______________ / ______________ / $______"),
            ("Property 2 address / lender / balance", "______________ / ______________ / $______"),
            ("Property 3 address / lender / balance", "______________ / ______________ / $______"),
            ("File path(s)", "______________"),
        ],
        "sample": "09_Mortgage_Statements_Cover_Sheet.md",
    },
    {
        "num": 10,
        "title": "Purchase Contract — 3887 Seneca",
        "source": "Standard CRE underwriting / deal file",
        "status": "ON FILE",
        "fields": [
            ("Contract date", "June 16, 2026"),
            ("Purchase price", "$500,000"),
            ("Seller concession", "$50,000"),
            ("Net consideration", "$450,000"),
            ("Deposit", "$10,000"),
            ("File path", "docs/contracts/3887 Seneca st Sale - Revised contract.pdf"),
        ],
        "sample": "10_Purchase_Contract_Cover_Sheet.md",
    },
    {
        "num": 11,
        "title": "Subject Property Pro Forma — Luxe Loft 716 / 3887 Seneca",
        "source": "Required for repayment analysis (not in IJB package body)",
        "status": "DRAFTED from Luxe deck — confirm rents",
        "fields": [
            ("9 salon suites annual", "$108,000"),
            ("Apt 1 / Apt 2 annual", "$13,200 / $15,600"),
            ("GPI", "$136,800"),
            ("NOI", "$95,920"),
            ("Debt service / DSCR (underwritten)", "$31,944 / 3.00"),
            ("File path", "docs/underwriting/SENECA_PROFORMA.md + Luxe PDF"),
        ],
        "sample": "11_Seneca_Pro_Forma_Template.md",
    },
    {
        "num": 12,
        "title": "Sources & Uses of Funds",
        "source": "Standard commercial loan request",
        "status": "PARTIAL — CapEx / loan ask open",
        "fields": [
            ("Net purchase (use)", "$450,000"),
            ("Closing costs (use)", "$______________"),
            ("CapEx / build-out (use)", "$______________"),
            ("Reserves (use)", "$______________"),
            ("Cornerstone loan (source)", "$______________"),
            ("Borrower equity (source)", "$______________"),
            ("File path", "docs/underwriting/SOURCES_AND_USES.md"),
        ],
        "sample": "12_Sources_and_Uses_Template.md",
    },
    {
        "num": 13,
        "title": "Entity Formation Documents — IJB Enterprises Inc.",
        "source": "Standard SPE borrower package",
        "status": "NEEDED",
        "fields": [
            ("Articles / Certificate of Incorporation", "Attached? Y/N — path: ______________"),
            ("EIN confirmation letter", "Attached? Y/N — path: ______________"),
            ("Operating agreement / bylaws", "Attached? Y/N — path: ______________"),
            ("Good standing / authority resolution", "Attached? Y/N — path: ______________"),
            ("Ownership % schedule", "______________"),
        ],
        "sample": "13_Entity_Documents_Cover_Sheet.md",
    },
    {
        "num": 14,
        "title": "Guarantor Tax Returns",
        "source": "Standard credit-union commercial underwriting",
        "status": "NEEDED",
        "fields": [
            ("Years requested", "Most recent ___ years (usually 2–3)"),
            ("Personal 1040s attached", "Y/N — path: ______________"),
            ("Business returns (if any) attached", "Y/N — path: ______________"),
            ("Extensions filed?", "Y/N"),
        ],
        "sample": "14_Tax_Returns_Cover_Sheet.md",
    },
    {
        "num": 15,
        "title": "Subject Property Rent Roll & Leases",
        "source": "Income verification for 3887 Seneca",
        "status": "NEEDED — resolve $2,400 vs $2,600/mo",
        "fields": [
            ("Apt 1 rent / lease end / tenant", "$______ / ________ / ________"),
            ("Apt 2 rent / lease end / tenant", "$______ / ________ / ________"),
            ("Combined monthly", "$______ (deck baseline $2,400)"),
            ("Security deposits held", "$______________"),
            ("Lease file paths", "______________"),
        ],
        "sample": "15_Rent_Roll_and_Leases_Template.md",
    },
    {
        "num": 16,
        "title": "CapEx / Contractor Bids — Suite Conversion",
        "source": "Build-out funding ask",
        "status": "NEEDED",
        "fields": [
            ("Partitions / doors", "$______________"),
            ("Electrical / data", "$______________"),
            ("Plumbing / wet stubs", "$______________"),
            ("HVAC / ventilation", "$______________"),
            ("Fire / life safety", "$______________"),
            ("Finishes", "$______________"),
            ("Soft costs + contingency", "$______________"),
            ("Total CapEx", "$______________"),
            ("Bid file paths", "______________"),
        ],
        "sample": "16_CapEx_Bid_Summary_Template.md",
    },
    {
        "num": 17,
        "title": "Floor Plans / Measured SF / Dental Architect Plans",
        "source": "Conversion feasibility (cited in Luxe deck)",
        "status": "NEEDED",
        "fields": [
            ("Measured 1st-floor SF (gross / rentable)", "______ / ______"),
            ("Suite count finalized", "______ (target 9)"),
            ("Architect plans attached?", "Y/N — path: ______________"),
            ("Sketch / layout attached?", "Y/N — path: ______________"),
        ],
        "sample": "17_Floor_Plan_Cover_Sheet.md",
    },
    {
        "num": 18,
        "title": "Insurance Quotes (Property + Liability)",
        "source": "Standard closing condition",
        "status": "NEEDED",
        "fields": [
            ("Carrier / agent", "______________"),
            ("Property coverage amount", "$______________"),
            ("Liability limits", "$______________"),
            ("Annual premium est.", "$______________"),
            ("Quote file path", "______________"),
        ],
        "sample": "18_Insurance_Quote_Cover_Sheet.md",
    },
    {
        "num": 19,
        "title": "Photo Package — Interior / Exterior",
        "source": "Underwriting / collateral review",
        "status": "NEEDED",
        "fields": [
            ("Exterior photos count", "______"),
            ("1st-floor / suite area photos", "______"),
            ("Apartment photos", "______"),
            ("Folder path", "______________"),
        ],
        "sample": "19_Photo_Package_Cover_Sheet.md",
    },
    {
        "num": 20,
        "title": "Government ID / Bio of Guarantor & Signers",
        "source": "KYC / credit-union CIP",
        "status": "NEEDED",
        "fields": [
            ("Driver license / passport copy", "Y/N — path: ______________"),
            ("Short bio / experience narrative", "Y/N — path: ______________"),
            ("SSN / DOB provided via secure channel?", "Y/N (do not email SSN in open mail)"),
        ],
        "sample": "20_ID_and_Bio_Cover_Sheet.md",
    },
    {
        "num": 21,
        "title": "Appraisal Order Authorization",
        "source": "Typical lender condition after terms",
        "status": "WHEN REQUESTED BY CORNERSTONE",
        "fields": [
            ("Borrower authorizes appraisal?", "Y/N"),
            ("Who pays appraisal fee?", "Borrower / Lender"),
            ("Preferred appraisal date window", "______________"),
            ("Contact for access", "______________"),
        ],
        "sample": "21_Appraisal_Authorization_Template.md",
    },
    {
        "num": 22,
        "title": "Loan Request Summary (Amount / Term / LTV)",
        "source": "Explicit ask missing from current package",
        "status": "FILL BEFORE SEND",
        "fields": [
            ("Requested loan amount", "$______________"),
            ("Requested term / amortization", "______ yrs / ______ yrs"),
            ("Requested rate type", "Fixed / Variable"),
            ("Requested LTV", "______ %"),
            ("Requested DSCR minimum understanding", "______ (deck shows 3.00)"),
            ("Guarantor(s)", "______________"),
        ],
        "sample": "22_Loan_Request_Summary_Template.md",
    },
]


def build_sample_templates():
    TEMPL.mkdir(parents=True, exist_ok=True)
    for item in ITEMS:
        path = TEMPL / item["sample"]
        lines = [
            f"# Item {item['num']}: {item['title']}",
            "",
            f"**Cornerstone package source:** {item['source']}  ",
            f"**Status:** {item['status']}  ",
            f"**Property / deal:** 3887 Seneca · Luxe Loft 716 · IJB Enterprises Inc.",
            "",
            "## Fill-in fields",
            "",
            "| Field | Value |",
            "|-------|-------|",
        ]
        for k, v in item["fields"]:
            lines.append(f"| {k} | {v} |")
        lines += [
            "",
            "## Sample narrative (edit freely)",
            "",
            f"This schedule supports Item {item['num']} of the Cornerstone Document Request Fill Form "
            f"for financing of 3887 Seneca Street, West Seneca, NY.",
            "",
            "Prepared by: ____________________  Date: __________",
            "",
            "Attachment checklist:",
            "- [ ] Primary document PDF/Excel attached",
            "- [ ] Cross-referenced in Supporting Documents index (Item 6)",
            "",
        ]
        # Extra sample tables for key templates
        if item["num"] == 8:
            lines += [
                "## Sample PFS line items",
                "",
                "| Assets | Amount | Liabilities | Amount |",
                "|--------|--------|-------------|--------|",
                "| Cash / banks | $____ | Credit cards | $____ |",
                "| Retirement | $____ | Auto loans | $____ |",
                "| Real estate (market) | $____ | Mortgages | $____ |",
                "| Other | $____ | Other | $____ |",
                "| **Total assets** | $____ | **Total liabilities** | $____ |",
                "",
                "**Net worth = Total assets − Total liabilities = $____**",
                "",
            ]
        if item["num"] == 11:
            lines += [
                "## Sample stabilized year",
                "",
                "| Income | Annual |",
                "|--------|--------|",
                "| Suite rents (9) | $108,000 |",
                "| Apt 1 | $13,200 |",
                "| Apt 2 | $15,600 |",
                "| **GPI** | **$136,800** |",
                "| − OpEx | ($40,880) |",
                "| **NOI** | **$95,920** |",
                "",
            ]
        if item["num"] == 12:
            lines += [
                "## Sample sources & uses grid",
                "",
                "| Uses | $ | Sources | $ |",
                "|------|---|--------|---|",
                "| Net purchase | 450,000 | Loan | ______ |",
                "| Closing | ______ | Equity | ______ |",
                "| CapEx | ______ | Deposit applied | 10,000 |",
                "| Reserves | ______ | Other | ______ |",
                "| **Total** | ______ | **Total** | ______ |",
                "",
            ]
        if item["num"] == 15:
            lines += [
                "## Sample rent roll",
                "",
                "| Unit | Tenant | Monthly rent | Lease start | Lease end | Deposit |",
                "|------|--------|--------------|-------------|-----------|---------|",
                "| Apt 1 | ______ | $____ | ______ | ______ | $____ |",
                "| Apt 2 | ______ | $____ | ______ | ______ | $____ |",
                "| **Total** | | $____ | | | $____ |",
                "",
            ]
        if item["num"] == 16:
            lines += [
                "## Sample CapEx bid log",
                "",
                "| Trade | Contractor | Bid $ | Date | Selected? |",
                "|-------|------------|-------|------|-----------|",
                "| Partitions | ______ | $____ | ______ | Y/N |",
                "| Electrical | ______ | $____ | ______ | Y/N |",
                "| Plumbing | ______ | $____ | ______ | Y/N |",
                "| HVAC | ______ | $____ | ______ | Y/N |",
                "| General / GC | ______ | $____ | ______ | Y/N |",
                "",
            ]
        path.write_text("\n".join(lines), encoding="utf-8")
    # master index of samples
    idx = TEMPL / "README.md"
    idx.write_text(
        "# Sample templates for Cornerstone document requests\n\n"
        "Each file matches a numbered item on `../Cornerstone_Document_Request_Fill_Form.docx`.\n\n"
        + "\n".join(f"- **Item {i['num']}:** [{i['sample']}]({i['sample']}) — {i['title']}" for i in ITEMS)
        + "\n",
        encoding="utf-8",
    )
    print("Wrote", len(ITEMS), "sample templates to", TEMPL)


def build_docx():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CORNERSTONE DOCUMENT REQUEST — FILL FORM")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x5C, 0x24, 0x30)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(
        "3887 Seneca Street · Luxe Loft 716 · IJB Enterprises Inc.\n"
        "Prepared for: Cornerstone Community Federal Credit Union — Michael Woods\n"
        "Based on: Commercial Financial Statement Package TOC + Supporting Documents references"
    )
    sr.font.size = Pt(10)

    n = doc.add_paragraph()
    nr = n.add_run(
        "Instructions: For each numbered point, fill every blank, mark Status, and attach the matching sample template "
        f"from docs/bank/sample_templates/. This completes Package Item 6 — Supporting Documents."
    )
    nr.italic = True
    nr.font.size = Pt(9)

    # summary tracker table
    h = doc.add_paragraph()
    hr = h.add_run("Master tracker")
    hr.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = RGBColor(0x5C, 0x24, 0x30)

    track = doc.add_table(rows=1 + len(ITEMS), cols=4)
    track.style = "Table Grid"
    hdr = ["#", "Document / Request", "Status (Done / In Progress / Needed)", "File path / notes"]
    for i, text in enumerate(hdr):
        track.rows[0].cells[i].text = text
        for p in track.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for idx, item in enumerate(ITEMS, start=1):
        row = track.rows[idx]
        row.cells[0].text = str(item["num"])
        row.cells[1].text = item["title"]
        row.cells[2].text = ""
        row.cells[3].text = ""
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()

    for item in ITEMS:
        ph = doc.add_paragraph()
        pr = ph.add_run(f"ITEM {item['num']}. {item['title']}")
        pr.bold = True
        pr.font.size = Pt(11)
        pr.font.color.rgb = RGBColor(0x5C, 0x24, 0x30)

        meta = doc.add_paragraph()
        mr = meta.add_run(f"Source: {item['source']}    |    Current status: {item['status']}    |    Sample template: sample_templates/{item['sample']}")
        mr.font.size = Pt(8)
        mr.italic = True

        table = doc.add_table(rows=len(item["fields"]) + 2, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Fill in / confirmed value"
        for p in table.rows[0].cells[0].paragraphs + table.rows[0].cells[1].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for i, (k, v) in enumerate(item["fields"], start=1):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        last = len(item["fields"]) + 1
        table.rows[last].cells[0].text = "My notes / date completed"
        table.rows[last].cells[1].text = "________________________________"
        doc.add_paragraph()

    cert = doc.add_paragraph()
    cr = cert.add_run("Certification")
    cr.bold = True
    cr.font.size = Pt(12)
    cr.font.color.rgb = RGBColor(0x5C, 0x24, 0x30)
    doc.add_paragraph(
        "I certify that the information and attachments referenced in this Document Request Fill Form are true and "
        "complete to the best of my knowledge, and I authorize Cornerstone Community Federal Credit Union to verify them."
    )
    sig = doc.add_table(rows=4, cols=2)
    sig.style = "Table Grid"
    sig_rows = [
        ("Print name", "Tiffany Durilla"),
        ("Entity", "IJB Enterprises Inc."),
        ("Signature", "________________________________"),
        ("Date", "________________"),
    ]
    for i, (k, v) in enumerate(sig_rows):
        sig.rows[i].cells[0].text = k
        sig.rows[i].cells[1].text = v

    doc.save(OUT_DOCX)
    print("Wrote", OUT_DOCX)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleR", parent=styles["Title"], fontSize=14, textColor=colors.HexColor("#5c2430"), leading=16))
    styles.add(ParagraphStyle("H", parent=styles["Heading2"], fontSize=10, textColor=colors.HexColor("#5c2430"), spaceBefore=8, spaceAfter=3))
    styles.add(ParagraphStyle("B", parent=styles["Normal"], fontSize=8, leading=10))
    styles.add(ParagraphStyle("I", parent=styles["Normal"], fontSize=7, leading=9, textColor=colors.HexColor("#555555")))

    story = []
    story.append(Paragraph("CORNERSTONE DOCUMENT REQUEST — FILL FORM", styles["TitleR"]))
    story.append(Paragraph(
        "3887 Seneca · Luxe Loft 716 · IJB Enterprises Inc.<br/>"
        "Cornerstone Community Federal Credit Union — Michael Woods<br/>"
        "Completes Package Item 6 — Supporting Documents",
        styles["B"],
    ))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#8b3a4a")))
    story.append(Paragraph(
        "<i>Fill every blank. Attach matching sample template from docs/bank/sample_templates/.</i>",
        styles["I"],
    ))

    # compact tracker
    story.append(Paragraph("Master tracker", styles["H"]))
    data = [[Paragraph("<b>#</b>", styles["B"]), Paragraph("<b>Document</b>", styles["B"]),
             Paragraph("<b>Status</b>", styles["B"]), Paragraph("<b>Path</b>", styles["B"])]]
    for item in ITEMS:
        data.append([
            Paragraph(str(item["num"]), styles["B"]),
            Paragraph(item["title"][:55], styles["B"]),
            Paragraph("☐ Done  ☐ IP  ☐ Need", styles["B"]),
            Paragraph("____________", styles["B"]),
        ])
    t = Table(data, colWidths=[0.35 * inch, 3.6 * inch, 1.5 * inch, 1.5 * inch])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f7f3ee")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(t)
    story.append(PageBreak())

    for item in ITEMS:
        block = []
        block.append(Paragraph(f"ITEM {item['num']}. {item['title']}", styles["H"]))
        block.append(Paragraph(f"Source: {item['source']} | Status: {item['status']} | Sample: {item['sample']}", styles["I"]))
        rows = [[Paragraph("<b>Field</b>", styles["B"]), Paragraph("<b>Fill in / value</b>", styles["B"])]]
        for k, v in item["fields"]:
            rows.append([Paragraph(k, styles["B"]), Paragraph(str(v), styles["B"])])
        rows.append([Paragraph("<b>Notes / date done</b>", styles["B"]), Paragraph("_______________________________", styles["B"])])
        tbl = Table(rows, colWidths=[2.4 * inch, 4.6 * inch])
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#cccccc")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#faf7f4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        block.append(tbl)
        block.append(Spacer(1, 6))
        story.append(KeepTogether(block))

    story.append(Paragraph("Certification", styles["H"]))
    story.append(Paragraph(
        "Information and attachments are true and complete to the best of my knowledge. Cornerstone may verify.",
        styles["B"],
    ))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Tiffany Durilla · IJB Enterprises Inc.", styles["B"]))
    story.append(Paragraph("Signature: ___________________________ &nbsp;&nbsp; Date: ______________", styles["B"]))

    SimpleDocTemplate(
        str(OUT_PDF), pagesize=letter,
        leftMargin=0.55 * inch, rightMargin=0.55 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
    ).build(story)
    print("Wrote", OUT_PDF)


if __name__ == "__main__":
    build_sample_templates()
    build_docx()
    build_pdf()
